from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "youtube_market_risk_indicator_series_script.docx"


def set_font(run, name="Malgun Gothic", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, color="000000", size=9.5):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)


def add_para(doc, text="", style=None, bold=False, size=None, color=None, after=5):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_code(doc, code):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    r = p.add_run(code)
    set_font(r, name="Consolas", size=8.5, color="1F2937")
    return p


def add_slide_table(doc, slides):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.55, 1.65, 2.05, 2.25]
    headers = ["#", "슬라이드", "화면 핵심 문장", "발표자 스크립트 요지"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width = Inches(widths[i])
        set_cell_shading(cell, "E8EEF5")
        set_cell_text(cell, h, bold=True, color="0B2545", size=8.7)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for idx, slide in enumerate(slides, 1):
        row = table.add_row()
        values = [str(idx), slide["title"], slide["screen"], slide["script"]]
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.width = Inches(widths[i])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_text(cell, value, size=8.3)
    return table


def add_slide_details(doc, slides):
    for idx, slide in enumerate(slides, 1):
        add_para(doc, f"Slide {idx}. {slide['title']}", style="Heading 3")
        add_para(doc, "화면 구성: " + slide["visual"], bold=True, size=10, color="1F4D78", after=3)
        add_para(doc, "발표 스크립트", bold=True, size=10, color="1F4D78", after=3)
        add_para(doc, slide["full"], size=10.2, after=7)
        if slide.get("demo"):
            add_para(doc, "간단한 데모", bold=True, size=10, color="1F4D78", after=3)
            add_para(doc, slide["demo"], size=10.2, after=5)
        if slide.get("code"):
            add_code(doc, slide["code"])


series = [
    {
        "name": "Series 1. 신용위험의 첫 번째 경고등: HY OAS와 HYG/LQD",
        "goal": "미국 주식 위험을 보기 전에, 먼저 기업 신용시장이 보내는 조기 경고를 읽는다.",
        "slides": [
            {
                "title": "주식시장 위험은 왜 채권시장에서 먼저 보일까",
                "screen": "주식은 기대를 가격에 반영하고, 신용시장은 생존 가능성을 가격에 반영한다.",
                "script": "주식보다 채권시장이 기업의 현금흐름 악화를 먼저 의심하는 경우가 많다는 점을 소개.",
                "visual": "좌측에는 S&P 500, 우측에는 하이일드 스프레드가 먼저 튀는 간단한 타임라인.",
                "full": "이번 시리즈의 출발점은 단순합니다. 주식시장이 화려하게 움직일 때도, 기업의 빚을 사는 투자자들은 훨씬 냉정하게 질문합니다. 이 회사가 이자를 계속 낼 수 있을까? 경기가 둔화되면 부도가 늘어날까? 그래서 신용시장은 종종 주식시장보다 먼저 위험을 말해줍니다. Series 1에서는 그중 가장 직관적인 두 지표, HY OAS와 HYG/LQD를 봅니다.",
            },
            {
                "title": "HY OAS란 무엇인가",
                "screen": "HY OAS = 하이일드 회사채가 국채보다 더 요구하는 보상.",
                "script": "OAS를 복잡한 옵션 수학이 아니라 위험 보상으로 설명.",
                "visual": "국채 수익률 위에 신용위험 프리미엄이 얹히는 계단형 그림.",
                "full": "HY OAS는 High Yield Option-Adjusted Spread입니다. 어렵게 들리지만 핵심은 간단합니다. 미국 하이일드 회사채에 투자하려면 안전한 국채보다 얼마나 더 많은 수익률을 요구하는가입니다. 이 추가 보상에는 부도위험, 유동성위험, 경기침체에 대한 공포가 섞여 있습니다. 숫자가 올라간다는 것은 투자자들이 위험한 회사채를 사기 위해 더 큰 보상을 요구한다는 뜻입니다.",
                "demo": "Dashboard의 Credit Risk 탭에서 HY_OAS z-score 라인을 보여주며 +1은 Watch, +2는 Danger 영역으로 설명합니다.",
            },
            {
                "title": "HY OAS가 상승할 때 시장은 무엇을 말하는가",
                "screen": "스프레드 확대는 위험 회피, 스프레드 축소는 위험 선호.",
                "script": "가격이 아니라 요구 보상이라는 관점에서 상승과 하락을 해석.",
                "visual": "HY OAS 상승 -> 채권가격 하락 -> 주식 위험 확대 가능성 순서도.",
                "full": "HY OAS가 높아질 때는 시장이 이렇게 말하는 것입니다. 같은 회사채라도 지금은 더 싼 가격, 더 높은 수익률이 아니면 사고 싶지 않다. 반대로 HY OAS가 낮아질 때는 위험을 감수하려는 자금이 돌아오고 있다는 뜻입니다. 단, 낮은 스프레드는 늘 좋은 신호만은 아닙니다. 너무 낮으면 시장이 위험을 과소평가하고 있을 수도 있습니다. 그래서 절대 수준과 변화 속도를 함께 봐야 합니다.",
            },
            {
                "title": "HYG와 LQD: 위험채권과 우량채권의 온도차",
                "screen": "HYG는 하이일드, LQD는 투자등급 회사채를 대표하는 ETF.",
                "script": "ETF 가격을 통해 신용위험 선호의 상대강도를 설명.",
                "visual": "HYG 박스에는 non-investment grade, LQD 박스에는 investment grade 라벨.",
                "full": "HYG는 미국 달러 표시 하이일드 회사채 ETF이고, LQD는 투자등급 회사채 ETF입니다. 둘 다 회사채지만 성격이 다릅니다. HYG는 경기와 부도위험에 더 민감하고, LQD는 상대적으로 우량하지만 금리 민감도가 더 큽니다. HYG/LQD 비율은 시장이 위험한 회사채를 우량 회사채보다 얼마나 선호하는지 보여주는 상대가격 지표입니다.",
            },
            {
                "title": "HYG/LQD 비율의 직관",
                "screen": "HYG/LQD 상승 = 위험선호 개선. 하락 = 신용위험 경계.",
                "script": "비율이 낮아지는 상황을 ‘위험자산에서 방어자산으로 이동’으로 설명.",
                "visual": "위험선호 바: HYG 쪽으로 기울면 risk-on, LQD 쪽으로 기울면 risk-off.",
                "full": "HYG/LQD 비율이 올라가면 하이일드가 투자등급보다 상대적으로 강하다는 뜻입니다. 투자자들이 부도위험을 감수하면서도 더 높은 수익을 찾고 있다는 신호입니다. 반대로 비율이 내려가면 하이일드가 더 약하다는 뜻이고, 시장은 위험한 신용을 줄이고 우량 신용을 선호하기 시작합니다. 이 프로젝트에서는 HYG/LQD의 z-score가 -1 아래로 내려가면 Watch, -2 아래로 내려가면 Danger로 해석합니다.",
            },
            {
                "title": "HY OAS와 HYG/LQD를 같이 봐야 하는 이유",
                "screen": "HY OAS는 스프레드, HYG/LQD는 시장 가격의 상대강도.",
                "script": "둘이 같은 방향으로 위험을 말할 때 신호 신뢰도가 높아진다는 메시지.",
                "visual": "2x2 매트릭스: OAS 상승/하락, HYG/LQD 상승/하락.",
                "full": "HY OAS와 HYG/LQD는 비슷해 보이지만 같은 지표가 아닙니다. HY OAS는 신용스프레드입니다. HYG/LQD는 ETF 가격으로 본 상대적 위험선호입니다. HY OAS가 오르고 HYG/LQD가 내려가면 둘 다 신용위험 확대를 말합니다. 이때는 단순한 노이즈보다 더 진지한 경고로 볼 수 있습니다. 반대로 둘이 엇갈리면 금리, ETF 수급, 단기 반등 같은 요인이 섞였는지 확인해야 합니다.",
            },
            {
                "title": "z-score로 위험을 표준화하는 이유",
                "screen": "서로 다른 단위의 지표를 ‘평소보다 얼마나 다른가’로 비교한다.",
                "script": "기술보다 해석 중심: 평균과 표준편차를 이용한 상대 위치 설명.",
                "visual": "정규분포 곡선 위에 +1, +2, -1, -2 구간 표시.",
                "full": "HY OAS는 퍼센트, HYG/LQD는 비율입니다. 단위가 다르기 때문에 원자료를 그대로 비교하면 해석이 흐려집니다. 그래서 z-score를 사용합니다. z-score는 오늘의 값이 최근 평균보다 얼마나 위나 아래에 있는지를 보여줍니다. HY OAS는 위로 튀면 위험이고, HYG/LQD는 아래로 꺾이면 위험입니다. 이 차이를 알고 봐야 차트가 훨씬 명확해집니다.",
                "code": "from indicators import add_credit_ratio, rolling_zscore_obs\n\nraw2 = add_credit_ratio(raw)\nhy_oas_z = rolling_zscore_obs(raw2['HY_OAS'], window_obs=60, min_obs=30)\nhyg_lqd_z = rolling_zscore_obs(raw2['HYG/LQD'], window_obs=126, min_obs=63)",
            },
            {
                "title": "간단한 데이터 흐름",
                "screen": "Yahoo Finance 가격 + FRED 신용스프레드 -> Credit Risk Dashboard.",
                "script": "Python은 지표를 설명하기 위한 도구로만 짧게 소개.",
                "visual": "Yahoo Finance, FRED, Python, Dash로 이어지는 4단계 파이프라인.",
                "full": "이 프로젝트에서 Python의 역할은 복잡한 모델을 자랑하는 것이 아닙니다. 필요한 데이터를 자동으로 불러오고, 같은 기준으로 정리하고, 대시보드에서 보기 쉽게 만드는 것입니다. HYG와 LQD는 Yahoo Finance 가격을 사용하고, HY OAS는 FRED의 ICE BofA US High Yield Index OAS를 사용합니다. 그다음 비율과 z-score를 계산해 Credit Risk 화면에서 함께 봅니다.",
                "code": "from data_fetching import fetch_data\nfrom indicators import add_credit_ratio\n\nraw = fetch_data(RISK_TICKERS, period='5y', include_treasury=True)\ncredit = add_credit_ratio(raw)[['HYG', 'LQD', 'HY_OAS', 'HYG/LQD']]",
            },
            {
                "title": "투자 판단에 바로 쓰기보다 경보판으로 쓰기",
                "screen": "신용지표는 매수/매도 버튼이 아니라 위험 온도계다.",
                "script": "단일 지표로 결론내리지 않고 다음 시리즈 지표와 결합한다고 연결.",
                "visual": "온도계: Normal, Watch, Danger 구간.",
                "full": "중요한 점은 HY OAS나 HYG/LQD 하나만 보고 매수와 매도를 결정하지 않는다는 것입니다. 이 지표들은 경보판입니다. 신용위험이 평소보다 커졌는지, 위험자산 선호가 약해졌는지를 먼저 알려줍니다. 다음 시리즈에서는 여기에 변동성, 유동성, 금리 스프레드를 더해 위험 신호가 정말 시장 전체로 번지고 있는지 확인하겠습니다.",
            },
        ],
    },
    {
        "name": "Series 2. 변동성, 유동성, 금리 스프레드: 위험이 시장 전체로 번지는 과정",
        "goal": "신용위험이 주식 변동성, 달러 유동성, 금리곡선과 어떻게 연결되는지 설명한다.",
        "slides": [
            {
                "title": "신용위험 다음에는 무엇을 확인할까",
                "screen": "신용이 흔들리면 변동성, 유동성, 금리곡선이 함께 반응한다.",
                "script": "Series 1의 신용 경고가 시장 전체 스트레스로 확산되는지 확인.",
                "visual": "Credit -> Volatility -> Liquidity -> Spread 화살표.",
                "full": "신용지표가 나빠졌다고 바로 시장 전체가 무너지는 것은 아닙니다. 그래서 다음 질문이 필요합니다. 위험이 신용시장 안에서만 머무는가, 아니면 주식 변동성, 달러 유동성, 금리곡선까지 번지는가? Series 2에서는 VIX 계열, 유동성 지표, 그리고 3개월-2년, 3개월-10년 스프레드를 함께 봅니다.",
            },
            {
                "title": "VIX: 보험료가 말하는 공포",
                "screen": "VIX는 향후 30일 S&P 500 변동성 기대를 반영한다.",
                "script": "VIX 20은 경계, 30 근처는 위험 구간으로 설명.",
                "visual": "VIX 차트 위 Watch 20, Danger 30 라인.",
                "full": "VIX는 흔히 공포지수라고 부릅니다. 정확히는 옵션시장이 반영하는 S&P 500의 향후 변동성 기대입니다. VIX가 20을 넘으면 평온한 시장과는 달라졌다고 보고, 30 근처로 가면 강한 위험회피 국면으로 해석할 수 있습니다. 이 프로젝트의 Volatility 탭도 VIX와 나스닥 변동성 지표인 VXN을 함께 보여줍니다.",
            },
            {
                "title": "VIX term structure: 단기 공포와 중기 공포",
                "screen": "단기 VIX가 중기 VIX보다 높아지면 시장은 당장의 충격을 가격에 반영한다.",
                "script": "VIX3M - VIX가 낮아지는 상황을 스트레스 신호로 설명.",
                "visual": "정상 곡선과 역전 곡선 비교.",
                "full": "보통은 먼 미래의 불확실성이 더 크기 때문에 중기 변동성이 단기 변동성보다 높습니다. 그런데 단기 VIX가 급등해 VIX3M과의 차이가 줄거나 역전되면 시장은 바로 눈앞의 충격을 가격에 반영하고 있다는 뜻입니다. 레짐 모델에서는 이 term structure도 위험환경 판단에 들어갑니다.",
            },
            {
                "title": "유동성: 시장을 움직이는 물의 양",
                "screen": "유동성이 풍부하면 위험자산이 버티고, 마르면 작은 충격도 커진다.",
                "script": "MMF, RRP, M2, 준비금 계열을 투자자 친화적으로 설명.",
                "visual": "저수지 비유: MMF/RRP/M2/Reserves가 수위에 기여.",
                "full": "유동성은 시장을 움직이는 물의 양처럼 볼 수 있습니다. 돈이 충분히 돌면 작은 악재가 흡수되지만, 유동성이 줄어든 시장에서는 같은 악재도 가격 충격이 커집니다. 이 프로젝트는 MMF, RRP, M2, 은행 준비금 및 준비금 프록시를 z-score로 바꾸어 글로벌 유동성 환경을 보여줍니다.",
                "demo": "Liquidity 탭에서 LiquidityScore와 각 구성요소가 같은 방향으로 움직이는지 확인합니다.",
            },
            {
                "title": "달러 강세 UUP: 글로벌 유동성의 압박",
                "screen": "달러가 강하면 글로벌 위험자산에는 유동성 압박이 커질 수 있다.",
                "script": "UUP를 달러 강세와 글로벌 금융조건 긴축의 proxy로 설명.",
                "visual": "달러 강세 -> 신흥국 부담 -> 위험자산 압박.",
                "full": "UUP는 달러 강세를 간단히 보는 ETF 지표입니다. 달러가 강해지면 미국 밖의 달러 부채 부담이 커지고, 글로벌 자금은 안전자산으로 이동하기 쉽습니다. 그래서 이 프로젝트에서는 UUP z-score가 높아지는 상황을 글로벌 유동성이 타이트해지는 신호로 해석합니다.",
            },
            {
                "title": "금리 스프레드: 경기 사이클의 압축 신호",
                "screen": "장단기 금리차 축소와 역전은 경기 둔화 기대를 반영한다.",
                "script": "3M-2Y, 3M-10Y 스프레드의 방향성과 해석 소개.",
                "visual": "정상 우상향 금리곡선과 역전 금리곡선 비교.",
                "full": "금리 스프레드는 채권시장이 보는 경기 사이클의 압축 신호입니다. 장기금리가 단기금리보다 낮아지는 역전은 시장이 앞으로 금리 인하와 경기 둔화를 예상한다는 뜻일 수 있습니다. 이 프로젝트는 2Y-3M, 10Y-3M, 10Y-2Y 스프레드를 만들어 레짐 점수에 반영합니다.",
            },
            {
                "title": "세 지표가 동시에 악화될 때",
                "screen": "신용 악화 + 변동성 상승 + 유동성 둔화 = 위험 신호의 합창.",
                "script": "한 지표보다 여러 지표의 동시성에 무게를 둔다.",
                "visual": "체크리스트: Credit, VIX, Liquidity, Spread.",
                "full": "시장 해석에서 중요한 것은 하나의 차트가 아니라 여러 지표의 동시성입니다. HY OAS가 상승하고 HYG/LQD가 하락하는데, VIX도 오르고 유동성 점수도 나빠지며 금리곡선도 경기 둔화를 말한다면 위험 신호의 신뢰도는 높아집니다. 이런 구조가 다음 시리즈의 글로벌 유동성, EEM 분석과 이어집니다.",
            },
            {
                "title": "Dashboard 데모: Stress Score로 묶어보기",
                "screen": "Stress Score는 여러 위험 지표를 하나의 온도계로 합친다.",
                "script": "VIX, HYG/LQD, 금리, UUP, EEM이 가중 결합된다는 점 설명.",
                "visual": "Stress Score 게이지 화면.",
                "full": "프로젝트의 Stress Score는 여러 시장 지표를 하나의 온도계로 합친 것입니다. VIX 계열, HYG/LQD, 금리, UUP, EEM이 각자의 방향으로 위험을 말하면 점수도 올라갑니다. 이 점수는 예측모델이라기보다 현재 시장환경을 빠르게 읽기 위한 대시보드 언어입니다.",
                "code": "z = compute_zscore(raw2, method='rolling', window_obs=252)\nstress = compute_stress_score(z)\n# Dashboard: Stress Score 탭에서 게이지와 추세를 확인",
            },
        ],
    },
    {
        "name": "Series 3. 글로벌 유동성과 EEM: 미국 밖에서 오는 위험 신호",
        "goal": "달러, 신흥국 주식, 글로벌 유동성의 관계를 설명한다.",
        "slides": [
            {
                "title": "왜 미국 주식 분석에 글로벌 지표가 필요할까",
                "screen": "미국 주식도 글로벌 달러 유동성의 영향을 받는다.",
                "script": "미국 주식시장이 세계 자금 흐름과 분리되어 있지 않다는 점 설명.",
                "visual": "미국 주식, 달러, 신흥국, 유동성 연결망.",
                "full": "미국 주식은 미국 기업의 가격처럼 보이지만, 실제로는 글로벌 자금의 중심에 있습니다. 달러가 강해지고 글로벌 유동성이 타이트해지면 신흥국과 위험자산이 먼저 흔들릴 수 있습니다. 그래서 Series 3에서는 UUP와 EEM을 통해 미국 밖에서 오는 위험 신호를 봅니다.",
            },
            {
                "title": "UUP: 달러 강세는 왜 위험 신호가 될 수 있나",
                "screen": "달러 강세는 글로벌 금융조건 긴축으로 작동할 수 있다.",
                "script": "달러 부채, 자금 회수, 위험자산 압박 설명.",
                "visual": "달러 상승 -> 해외 달러부채 부담 -> 위험자산 매도 압력.",
                "full": "달러 강세는 단순한 환율 문제가 아닙니다. 글로벌 금융시장에서 달러는 결제통화이자 부채통화입니다. 달러가 강해지면 달러를 빌린 국가와 기업의 부담이 커지고, 자금은 위험자산에서 안전자산으로 이동하기 쉬워집니다. 그래서 UUP의 상승은 글로벌 유동성 압박을 보여주는 간단한 proxy가 됩니다.",
            },
            {
                "title": "EEM: 신흥국은 위험 선호의 민감한 센서",
                "screen": "EEM 약세는 글로벌 위험회피가 커지고 있다는 신호일 수 있다.",
                "script": "신흥국 주식은 달러, 원자재, 글로벌 성장 기대에 민감하다고 설명.",
                "visual": "EEM 가격과 6개월 고점 대비 drawdown 표시.",
                "full": "EEM은 신흥국 주식 ETF입니다. 신흥국은 달러 강세, 글로벌 성장 둔화, 자금 유출에 민감하게 반응합니다. 그래서 EEM이 6개월 고점 대비 크게 빠지고 있다면, 글로벌 투자자들이 위험을 줄이고 있다는 신호로 볼 수 있습니다. 프로젝트에서는 EEM drawdown을 계산해 최근 3개월 변화도 함께 보여줍니다.",
            },
            {
                "title": "UUP와 EEM을 함께 보는 법",
                "screen": "UUP 상승 + EEM 하락은 글로벌 risk-off 조합.",
                "script": "둘의 반대 방향 움직임이 글로벌 유동성 악화의 강한 표현이라고 설명.",
                "visual": "2축 차트 또는 대각선 매트릭스.",
                "full": "UUP가 상승하고 EEM이 하락하면 글로벌 risk-off 조합입니다. 달러는 강해지고 신흥국 주식은 약해지는 환경입니다. 이때 미국 대형주가 당장은 버티더라도 시장 내부의 위험 선호는 이미 약해지고 있을 수 있습니다. 반대로 UUP가 약해지고 EEM이 회복되면 글로벌 위험 선호가 개선되는 신호로 볼 수 있습니다.",
            },
            {
                "title": "글로벌 유동성 점수의 역할",
                "screen": "여러 유동성 지표를 하나로 묶으면 방향성이 보인다.",
                "script": "MMF, RRP, M2, 준비금 프록시와 credit adjustment 소개.",
                "visual": "LiquidityScore 구성요소 막대.",
                "full": "개별 유동성 지표는 발표 주기와 의미가 다릅니다. MMF, RRP, M2, 준비금 계열을 각각 보면 복잡하지만, z-score로 표준화하고 합치면 전체 유동성의 방향성을 볼 수 있습니다. 레짐 모델에서는 여기에 HY OAS와 HYG/LQD도 반영해, 유동성이 좋아 보이더라도 신용위험이 커지는 상황을 보정합니다.",
            },
            {
                "title": "Dashboard 데모: Global Liquidity와 EEM",
                "screen": "UUP z-score, EEM drawdown, LiquidityScore를 함께 확인한다.",
                "script": "기술은 짧게, 화면에서 해석하는 순서 강조.",
                "visual": "Global Liquidity 탭과 Global Risk (EEM) 탭 전환.",
                "full": "대시보드에서는 먼저 Global Liquidity 탭에서 UUP와 유동성 지표를 확인합니다. 다음으로 Global Risk (EEM) 탭에서 EEM이 6개월 고점 대비 얼마나 빠져 있는지 봅니다. 마지막으로 이 두 신호가 Series 1의 신용위험, Series 2의 변동성과 같은 방향인지 확인합니다. 방향이 같을수록 위험 신호는 강해집니다.",
                "code": "eem_dd = (raw['EEM'] / raw['EEM'].rolling(126).max() - 1) * 100\nuup_z = (raw['UUP'] - raw['UUP'].mean()) / raw['UUP'].std()",
            },
            {
                "title": "글로벌 신호의 한계",
                "screen": "EEM은 신호이지 정답이 아니다.",
                "script": "중국 비중, 환율, 지역 이슈 등 노이즈를 설명.",
                "visual": "EEM 해석 체크리스트.",
                "full": "EEM은 좋은 글로벌 위험 센서지만 완벽하지는 않습니다. 특정 국가 비중, 환율, 원자재 가격, 지역 정치 리스크가 섞일 수 있습니다. 그래서 EEM은 단독 결론이 아니라 글로벌 위험 선호를 확인하는 보조지표로 쓰는 것이 좋습니다. 핵심은 여러 지표가 같은 이야기를 하는지 보는 것입니다.",
            },
        ],
    },
    {
        "name": "Series 4. 레짐 모니터와 검증: 시장 상태를 분류하고 맞았는지 확인하기",
        "goal": "risk_on부터 crisis까지의 레짐 점수와 정확성 검증 구조를 설명한다.",
        "slides": [
            {
                "title": "레짐이란 무엇인가",
                "screen": "레짐은 시장의 현재 날씨를 다섯 단계로 분류하는 언어다.",
                "script": "risk_on, neutral, caution, risk_off, crisis 소개.",
                "visual": "초록에서 진한 빨강으로 이어지는 5단계 바.",
                "full": "레짐은 시장 상태를 설명하는 언어입니다. 이 프로젝트는 시장을 risk_on, neutral, caution, risk_off, crisis 다섯 단계로 나눕니다. 중요한 것은 레짐이 미래를 완벽히 맞히는 점쟁이가 아니라, 현재 위험환경을 일관된 기준으로 읽는 도구라는 점입니다.",
            },
            {
                "title": "레짐 점수의 구성",
                "screen": "유동성, 신용, 변동성, 달러, EEM, 금리 스프레드를 하나로 합친다.",
                "script": "가중치 방향: 유동성/EEM은 플러스, OAS/VIX/UUP/Stress는 마이너스.",
                "visual": "긍정 기여와 부정 기여를 나눈 저울 그림.",
                "full": "레짐 점수는 여러 지표의 가중합입니다. 유동성 점수와 EEM은 좋아질수록 레짐을 개선합니다. HY OAS, VIX, Stress Score, UUP는 높아질수록 레짐을 악화시킵니다. HYG/LQD는 높아질수록 위험 선호가 좋다고 봅니다. 즉 한 지표가 아니라 시장의 여러 층을 하나의 상태 점수로 요약합니다.",
            },
            {
                "title": "confidence: 지표들이 같은 말을 하는가",
                "screen": "신뢰도는 지표 간 합의 정도와 데이터 커버리지를 반영한다.",
                "script": "방향을 맞춘 지표들의 분산이 낮을수록 confidence가 높다고 설명.",
                "visual": "여러 화살표가 같은 방향이면 high confidence, 흩어지면 low confidence.",
                "full": "레짐 라벨만 보는 것보다 confidence를 함께 보는 것이 중요합니다. 여러 지표가 같은 방향을 말하면 신뢰도가 높아집니다. 반대로 신용은 나쁜데 유동성은 좋고, VIX는 낮은데 EEM은 약하면 신뢰도는 낮아질 수 있습니다. 이때는 시장이 전환점에 있거나 지표 간 시차가 생긴 상태일 수 있습니다.",
            },
            {
                "title": "레짐 정확성은 어떻게 확인하나",
                "screen": "각 레짐 이후 21거래일 SPY 수익률이 기대 방향과 맞았는지 본다.",
                "script": "risk_on/neutral/caution은 상승 기대, risk_off/crisis는 하락 또는 방어 기대.",
                "visual": "Regime -> 21D Forward Return -> Hit/Miss 흐름.",
                "full": "레짐 모니터는 반드시 검증이 필요합니다. 이 프로젝트는 스냅샷 이력과 SPY 가격을 결합해 각 레짐 이후 21거래일 수익률을 확인합니다. risk_on, neutral, caution은 이후 SPY가 오르면 hit로 보고, risk_off와 crisis는 이후 SPY가 flat 또는 하락하면 hit로 봅니다. 완벽한 예측보다 방향성이 실제로 도움이 되는지 확인하는 과정입니다.",
            },
            {
                "title": "Stress vs Forward Returns의 다른 질문",
                "screen": "레짐 정확성과 스트레스 이후 반등 분석은 서로 다른 질문이다.",
                "script": "고스트레스 이후 상승은 반등 분석에서는 positive지만 레짐 방향성 검증과 다름.",
                "visual": "두 질문 비교: Regime Accuracy vs Rebound Analysis.",
                "full": "Historical Regime Accuracy와 Stress vs Forward Returns는 비슷해 보이지만 질문이 다릅니다. 레짐 정확성은 해당 레짐이 말한 방향이 맞았는지를 봅니다. 반면 Stress vs Forward Returns는 스트레스가 높았던 날 이후 SPY가 반등했는지를 봅니다. 위기 국면에서 이후 반등이 나오면 스트레스 분석에서는 긍정적 결과지만, 레짐 방향성 검증에서는 다른 의미를 가질 수 있습니다.",
            },
            {
                "title": "중복 신호를 에피소드로 묶는 이유",
                "screen": "위기 때 매일 나온 신호를 하나의 독립 사건으로 정리한다.",
                "script": "21거래일 안의 신호를 묶어 과대평가를 피하는 구조 설명.",
                "visual": "여러 신호 점을 하나의 episode 박스로 묶는 타임라인.",
                "full": "위기 때는 스트레스 신호가 여러 날 연속으로 나옵니다. 이 모든 날을 독립 표본처럼 계산하면 같은 사건을 여러 번 세는 문제가 생깁니다. 그래서 프로젝트는 21거래일 안에 발생한 crisis 신호를 하나의 에피소드로 묶고, 그중 스트레스가 가장 높았던 날을 대표 신호로 사용합니다. 이 방식이 과도한 자신감을 줄입니다.",
            },
            {
                "title": "Dashboard 데모: Regime Monitor",
                "screen": "현재 레짐, 점수, 신뢰도, 전환 경고를 한 화면에서 본다.",
                "script": "레짐은 매일의 시장 브리핑 도구로 사용하는 방식을 설명.",
                "visual": "Regime Monitor 카드, regime_score 추세, regime_state 차트.",
                "full": "Regime Monitor 탭에서는 현재 레짐 라벨, 레짐 점수, confidence, 거래 허용 여부, 포지션 크기 배수를 확인할 수 있습니다. 이것은 매일의 시장 브리핑 도구입니다. 오늘 시장은 공격적인가, 중립인가, 방어적인가? 그리고 그 판단을 여러 지표가 얼마나 강하게 지지하고 있는가를 한눈에 보는 화면입니다.",
            },
        ],
    },
    {
        "name": "Series 5. Stress Score로 결론내기: 위험 온도계를 투자 프로세스에 넣는 법",
        "goal": "여러 지표를 종합해 실전적인 위험 점검 루틴으로 마무리한다.",
        "slides": [
            {
                "title": "왜 최종 결론은 Stress Score인가",
                "screen": "여러 지표를 하나의 위험 온도계로 압축한다.",
                "script": "개별 지표의 노이즈를 줄이고 매일 체크할 수 있는 점수로 만든다는 메시지.",
                "visual": "여러 지표가 하나의 게이지로 모이는 그림.",
                "full": "지표가 많아지면 오히려 판단이 어려워질 수 있습니다. Stress Score는 여러 위험 지표를 하나의 온도계로 압축합니다. VIX 계열, HYG/LQD, 금리, UUP, EEM이 함께 반영되기 때문에 개별 지표의 노이즈를 줄이고 시장 스트레스의 큰 방향을 보기에 좋습니다.",
            },
            {
                "title": "Stress Score 버킷 해석",
                "screen": "Very Easy, Easy, Normal, Stress, Crisis로 구간을 나눈다.",
                "script": "점수별 의미와 포지션 조절 아이디어 소개.",
                "visual": "5단계 색상 바와 점수 구간.",
                "full": "프로젝트의 스트레스 버킷은 Very Easy, Easy, Normal, Stress, Crisis로 나뉩니다. Very Easy와 Easy는 위험자산 환경이 상대적으로 편안한 구간입니다. Normal은 중립, Stress는 경계, Crisis는 강한 위험회피 구간입니다. 이 구분은 미래를 보장하지 않지만, 투자자가 자신의 포지션 크기와 리스크 예산을 조절하는 기준점이 됩니다.",
            },
            {
                "title": "Stress Score와 레짐을 함께 쓰기",
                "screen": "점수는 온도, 레짐은 날씨 설명이다.",
                "script": "Stress Score >= 2와 crisis 레짐의 결합이 강한 위기 조건이라고 설명.",
                "visual": "Stress Score 게이지 + Regime 라벨 카드.",
                "full": "Stress Score는 시장의 온도입니다. 레짐은 그 온도를 포함한 전체 날씨 설명입니다. 이 프로젝트에서는 Stress Score가 2 이상이고 레짐이 crisis일 때를 주요 위기 조건으로 봅니다. 단순히 점수만 높은 날보다, 여러 지표가 동시에 위기를 말하는 날을 더 강한 신호로 보는 방식입니다.",
            },
            {
                "title": "매일 5분 위험 점검 루틴",
                "screen": "Credit -> Volatility -> Liquidity -> Global -> Regime -> Stress.",
                "script": "시리즈 전체를 하나의 실전 루틴으로 정리.",
                "visual": "6단계 체크리스트.",
                "full": "매일의 루틴은 간단합니다. 첫째, HY OAS와 HYG/LQD로 신용위험을 봅니다. 둘째, VIX와 VIX term structure로 변동성을 봅니다. 셋째, 유동성 점수와 UUP를 확인합니다. 넷째, EEM으로 글로벌 위험 선호를 봅니다. 다섯째, Regime Monitor에서 종합 상태와 confidence를 확인합니다. 마지막으로 Stress Score가 어느 버킷에 있는지 봅니다.",
            },
            {
                "title": "실전 해석 예시",
                "screen": "경고는 빨리, 결론은 천천히.",
                "script": "신용 먼저 악화, 변동성 확인, 레짐 변화까지 단계적으로 판단.",
                "visual": "Day 1, Day 3, Day 10 식의 신호 누적 타임라인.",
                "full": "예를 들어 HY OAS가 먼저 상승하고 HYG/LQD가 하락하기 시작했다고 합시다. 이때 바로 모든 포지션을 정리하는 것이 아니라 경고등을 켭니다. 이후 VIX가 상승하고, 유동성 점수가 악화되고, EEM이 약해지며 레짐이 caution에서 risk_off로 내려가면 신호가 누적됩니다. 경고는 빨리 받고, 결론은 여러 지표가 확인해 줄 때 내리는 방식입니다.",
            },
            {
                "title": "Python은 자동화 도구일 뿐이다",
                "screen": "핵심은 코드가 아니라 지표의 의미와 해석 규칙이다.",
                "script": "기술 중심 채널이 아니라 투자지표 교육 채널이라는 방향 강조.",
                "visual": "데이터 수집, 계산, Dashboard, 해석 네 단계.",
                "full": "이 시리즈에서 Python은 주인공이 아닙니다. 주인공은 시장위험 지표입니다. Python은 데이터를 자동으로 가져오고, z-score를 계산하고, Dash로 화면을 만들어주는 도구입니다. 투자자에게 중요한 것은 코드 한 줄보다, HY OAS가 왜 오르는지, HYG/LQD가 왜 내려가는지, VIX와 EEM이 같은 방향을 말할 때 무엇을 조심해야 하는지입니다.",
            },
            {
                "title": "시리즈 결론",
                "screen": "위험을 맞히려 하지 말고, 위험이 커지는 구조를 관찰하라.",
                "script": "채널 전체의 메시지로 마무리.",
                "visual": "전체 지표 맵과 Stress Score 게이지.",
                "full": "시장위험 분석의 목적은 내일의 수익률을 정확히 맞히는 것이 아닙니다. 위험이 커지는 구조를 남들보다 일관되게 관찰하는 것입니다. 신용위험, 변동성, 유동성, 글로벌 위험, 레짐, 스트레스 점수를 하나의 프로세스로 묶으면 시장을 훨씬 차분하게 볼 수 있습니다. 이 시리즈의 결론은 이것입니다. 예측보다 중요한 것은 반복 가능한 위험 점검 체계입니다.",
            },
        ],
    },
]


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.75)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)

    styles = doc.styles
    for style_name in ["Normal", "Heading 1", "Heading 2", "Heading 3"]:
        style = styles[style_name]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")

    styles["Normal"].font.size = Pt(10.5)
    styles["Heading 1"].font.size = Pt(17)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string("0B2545")
    styles["Heading 2"].font.size = Pt(13.5)
    styles["Heading 2"].font.color.rgb = RGBColor.from_string("2E74B5")
    styles["Heading 3"].font.size = Pt(11.5)
    styles["Heading 3"].font.color.rgb = RGBColor.from_string("1F4D78")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    r = title.add_run("미국 주식시장 위험지표 YouTube 시리즈")
    set_font(r, size=22, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run("PowerPoint 제작용 슬라이드 구성 및 발표 스크립트")
    set_font(sr, size=12, color="555555")

    add_para(doc, "문서 목적", style="Heading 1")
    add_para(
        doc,
        "이 문서는 yfinance-risk-analysis 프로젝트의 지표 구조를 바탕으로, 기술 구현보다 미국 주식시장 위험을 읽는 지식과 해석에 초점을 둔 YouTube 시리즈용 PPT 스크립트입니다. 각 슬라이드는 화면에 들어갈 문장과 발표자 노트를 분리해 바로 PowerPoint로 옮기기 쉽게 구성했습니다.",
        size=10.5,
    )
    add_para(doc, "권장 시리즈 구성", style="Heading 1")
    overview = doc.add_table(rows=1, cols=3)
    overview.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["시리즈", "핵심 질문", "주요 지표"]):
        set_cell_shading(overview.rows[0].cells[i], "E8EEF5")
        set_cell_text(overview.rows[0].cells[i], h, bold=True, color="0B2545", size=9)
    rows = [
        ("1", "신용시장은 주식시장 위험을 어떻게 먼저 말하는가?", "HY OAS, HYG/LQD, HYG, LQD"),
        ("2", "위험이 시장 전체로 번지고 있는가?", "VIX, VXN, VIX term, LiquidityScore, 금리 스프레드"),
        ("3", "미국 밖의 위험 선호는 어떤 신호를 보내는가?", "UUP, EEM, 글로벌 유동성"),
        ("4", "여러 지표를 하나의 시장 상태로 분류할 수 있는가?", "Regime Score, Confidence, Regime Accuracy"),
        ("5", "최종적으로 위험 점검 루틴을 어떻게 만들 것인가?", "Stress Score, Stress Buckets, Crisis Episodes"),
    ]
    for row_values in rows:
        row = overview.add_row()
        for i, value in enumerate(row_values):
            set_cell_text(row.cells[i], value, size=8.7)

    add_para(doc, "공통 톤", style="Heading 1")
    add_para(doc, "투자 조언이 아니라 시장위험을 읽는 교육 콘텐츠로 진행합니다. 기술 구현은 데이터 수집과 대시보드 확인 정도로 짧게 보여주고, 대부분의 시간은 지표의 의미, 상승/하락 해석, 여러 지표가 동시에 움직일 때의 판단법에 사용합니다.", size=10.5)

    for idx, s in enumerate(series, 1):
        doc.add_section(WD_SECTION.NEW_PAGE)
        add_para(doc, s["name"], style="Heading 1")
        add_para(doc, "목표: " + s["goal"], bold=True, size=10.5, color="1F4D78")
        add_para(doc, "슬라이드 요약", style="Heading 2")
        add_slide_table(doc, s["slides"])
        add_para(doc, "상세 발표 스크립트", style="Heading 2")
        add_slide_details(doc, s["slides"])

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_para(doc, "프로젝트와 연결되는 핵심 파일", style="Heading 1")
    add_para(doc, "app.py: 대시보드 탭과 주요 티커 그룹 정의", size=10.2)
    add_para(doc, "data_fetching.py: Yahoo Finance와 FRED 데이터 수집, HY_OAS 포함", size=10.2)
    add_para(doc, "indicators.py: HYG/LQD 비율, z-score, Stress Score 계산", size=10.2)
    add_para(doc, "regime.py: LiquidityScore, Regime Score, confidence, 레짐 라벨 생성", size=10.2)
    add_para(doc, "research_metrics.py: 레짐 정확성, Stress vs Forward Returns, crisis episode 분석", size=10.2)

    add_para(doc, "출처 참고", style="Heading 1")
    sources = [
        "FRED: ICE BofA US High Yield Index Option-Adjusted Spread (BAMLH0A0HYM2), https://fred.stlouisfed.org/series/BAMLH0A0HYM2",
        "BlackRock iShares: HYG, iShares iBoxx $ High Yield Corporate Bond ETF, https://www.blackrock.com/us/individual/products/239565/ishares-iboxx-high-yield-corporate-bond-etf",
        "BlackRock iShares: LQD, iShares iBoxx $ Investment Grade Corporate Bond ETF, https://www.blackrock.com/us/individual/products/239566/ishares-iboxx-investment-grade-corporate-bond-etf",
    ]
    for source in sources:
        add_para(doc, source, size=9.5)

    add_para(doc, "주의 문구", style="Heading 1")
    add_para(doc, "본 콘텐츠는 투자교육 목적이며 특정 자산의 매수, 매도, 보유 권유가 아닙니다. 모든 지표는 과거 및 현재 시장 데이터를 해석하는 도구이며, 미래 수익률을 보장하지 않습니다.", size=10.2)

    doc.save(OUT)


if __name__ == "__main__":
    build()
